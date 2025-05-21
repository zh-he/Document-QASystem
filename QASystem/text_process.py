import hashlib
import re
from pathlib import Path
import fitz  # PyMuPDF
import unicodedata
import logging
import docx  # python-docx for Word files
from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

logger = logging.getLogger(__name__)


class TextProcessor:
    """统一处理PDF、TXT和Word文件的加载和文本清洗"""

    @staticmethod
    def _load_pdf_raw(file_path: Path) -> str:
        """从PDF提取原始文本"""
        try:
            with fitz.open(str(file_path)) as doc:
                # 使用更精确的块提取，尝试保留段落结构
                full_text = []
                for page in doc:
                    blocks = page.get_text("dict", flags=fitz.TEXTFLAGS_TEXT).get("blocks", [])
                    for block in blocks:
                        if block['type'] == 0:  # 文本块
                            block_text = ""
                            for line in block.get("lines", []):
                                line_text = "".join(span['text'] for span in line.get("spans", []))
                                block_text += line_text + "\n"  # 行末加换行符
                            if block_text.strip():
                                full_text.append(block_text.strip())
                # 使用单个换行符连接块，因为 PyMuPDF 的块通常对应段落
                text = "\n".join(full_text)
                return text
        except Exception as e:
            logger.error(f"PDF加载失败 {file_path.name}: {str(e)}")
            return ""

    @staticmethod
    def _load_txt_raw(file_path: Path) -> str:
        """从TXT文件提取原始文本"""
        try:
            # 尝试使用不同编码读取文本文件
            encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                        return text
                except UnicodeDecodeError:
                    continue

            # 如果所有编码都失败，尝试二进制读取并解码
            logger.warning(f"无法使用常见编码读取{file_path.name}，尝试二进制读取")
            with open(file_path, 'rb') as file:
                raw_bytes = file.read()
                return raw_bytes.decode('utf-8', errors='replace')
        except Exception as e:
            logger.error(f"TXT文件加载失败 {file_path.name}: {str(e)}")
            return ""

    @staticmethod
    def _load_docx_raw(file_path: Path) -> str:
        """从Word文档提取原始文本"""
        try:
            doc = docx.Document(file_path)
            # 保留段落结构，每个段落后添加换行符
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n".join(paragraphs)
        except Exception as e:
            logger.error(f"Word文档加载失败 {file_path.name}: {str(e)}")
            return ""

    @staticmethod
    def apply_custom_regex(text: str, custom_regex: str) -> str:
        """应用自定义的多行正则表达式进行替换"""
        if not custom_regex:
            return text
        try:
            regex_lines = custom_regex.strip().split('\n')
            processed_text = text
            for line in regex_lines:
                pattern = line.strip()
                if pattern:
                    # 假设自定义正则也是用来删除匹配项
                    processed_text = re.sub(pattern, '', processed_text)
            return processed_text.strip()
        except re.error as e:
            logger.error(f"自定义正则表达式错误: {e}，正则: '{pattern}'")
            # 根据需要决定是返回原文还是抛出异常
            return text  # 或者 raise ValueError(f"自定义正则表达式错误: {e}") from e

    @staticmethod
    def clean_text(text: str) -> str:
        """
        对文本进行通用清洗：
        1. 标准化空白字符
        2. 移除特殊控制字符
        3. 合并多余空行
        """
        if not text:
            return ""

        # 标准化空白字符：连续空格替换为单个空格
        text = re.sub(r'[ \t]+', ' ', text)

        # 合并多余空行（3个以上换行符变为2个）
        text = re.sub(r'\n{3,}', '\n\n', text)

        # 确保文件开头和结尾无多余空白
        return text.strip()

    @classmethod
    def extract_and_process_text(cls, file_path: Path, custom_regex: str | None = None) -> str:
        """
        根据文件类型加载文件，然后根据是否有自定义正则，选择应用自定义规则或默认清洗规则。
        """
        # 确保file_path是Path对象
        if isinstance(file_path, str):
            file_path = Path(file_path)

        # 根据文件扩展名判断文件类型并加载
        suffix = file_path.suffix.lower()
        raw_text = ""

        if suffix == '.pdf':
            raw_text = cls._load_pdf_raw(file_path)
        elif suffix == '.txt':
            raw_text = cls._load_txt_raw(file_path)
        elif suffix in ['.docx', '.doc']:
            raw_text = cls._load_docx_raw(file_path)
        else:
            logger.warning(f"不支持的文件类型: {suffix}，文件: {file_path.name}")
            return ""

        if not raw_text:
            return ""

        # 应用文本清洗
        cleaned_text = cls.clean_text(raw_text)

        # 应用自定义正则（如果有）
        if custom_regex:
            processed_text = cls.apply_custom_regex(cleaned_text, custom_regex)
            logger.info(f"从 {file_path} 提取并应用自定义正则处理完成。")
        else:
            processed_text = cleaned_text
            logger.info(f"从 {file_path} 提取并应用标准清洗处理完成。")

        return processed_text


class DocumentSplitter:
    """封装文本分割逻辑"""

    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 100):
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            keep_separator=False
        )

    def split_text(self, text: str, metadata: dict) -> list[Document]:
        """将单个文本和元数据分割成 Document 列表"""
        text_chunks = self.splitter.split_text(text)
        documents = [Document(page_content=chunk, metadata=metadata.copy()) for chunk in text_chunks]
        return documents

    def split_documents(self, docs: list[Document]) -> list[Document]:
        """对已有的 Document 列表进行分割 (通常用于初始化)"""
        return self.splitter.split_documents(docs)


def hash_filename(filename: str) -> str:
    """将文件名进行 MD5 哈希处理，返回哈希值字符串"""
    return hashlib.md5(filename.encode('utf-8')).hexdigest()
